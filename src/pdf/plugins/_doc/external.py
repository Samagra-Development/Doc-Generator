import os
import traceback
import aspose.words as aw
from docx import Document

from ...plugins._pdf.external import PDFPlugin
from ...uploaders.minio import MinioUploader
from ...utils import build_doc


def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None


class DOCXPlugin(PDFPlugin):
    """
        Plugin class which extend from PDFPlugin
    """

    def build_file(self, template):
        """
        Function to build PDF and return a file (fetch template and build pdf)
        """
        is_successful = error_code = error_msg = None
        drive_file_loc = f'pdf/drivefiles/{self.token}.docx'
        try:
            doc = aw.Document()
            builder = aw.DocumentBuilder(doc);
            builder.insert_html(template)
            doc.save(drive_file_loc)
            document = Document(drive_file_loc)
            paragraph = document.paragraphs[0]
            print(paragraph.text)
            delete_paragraph(paragraph)
            document.save(drive_file_loc)
            is_successful = True
        except Exception as e:
            traceback.print_exc()
            error_msg = f"Failed to generate doc: {e}"
            error_code = 803
        return error_code, error_msg, is_successful

    def upload_file(self):
        """
        Function to save PDF
        """
        error_code = error_msg = final_data = None
        try:
            drive_file_loc = f'pdf/drivefiles/{self.token}.docx'
            if self.uploader == "minio":
                host = self.user_config["MINIO_HOST"]
                access_key = self.user_config["MINIO_ACCESS_KEY"]
                secret_key = self.user_config["MINIO_SECRET_KEY"]
                bucket_name = self.user_config["MINIO_BUCKET_NAME"]
                uploader = MinioUploader(host, access_key, secret_key, bucket_name)
                error_code, error_msg, final_data = uploader.put(f'{self.token}.docx', f'{self.token}.docx', None)
                if error_code is None:
                    if os.path.exists(drive_file_loc):
                        os.remove(drive_file_loc)
                else:
                    raise Exception("Failed to build the pdf")
                return error_code, error_msg, final_data
            else:
                raise Exception("Uploader plugin not supported")
        except Exception as e:
            traceback.print_exc()
            error_code = 805
            error_msg = f"Something went wrong: {e}"
        finally:
            return error_code, error_msg, final_data